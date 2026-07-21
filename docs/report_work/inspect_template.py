from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def length(value):
    return None if value is None else round(value.pt, 2)


def font_name(run):
    rpr = run._element.rPr
    fonts = None if rpr is None else rpr.rFonts
    return {
        "name": run.font.name,
        "ascii": None if fonts is None else fonts.get(qn("w:ascii")),
        "eastAsia": None if fonts is None else fonts.get(qn("w:eastAsia")),
        "hAnsi": None if fonts is None else fonts.get(qn("w:hAnsi")),
    }


def para_info(index, p):
    fmt = p.paragraph_format
    runs = []
    for run in p.runs:
        if not run.text.strip():
            continue
        color = run.font.color.rgb
        runs.append(
            {
                "text": run.text[:100],
                "font": font_name(run),
                "size_pt": length(run.font.size),
                "bold": run.bold,
                "italic": run.italic,
                "color": None if color is None else str(color),
            }
        )
    return {
        "index": index,
        "text": p.text[:220],
        "style": p.style.name,
        "alignment": None if p.alignment is None else int(p.alignment),
        "space_before_pt": length(fmt.space_before),
        "space_after_pt": length(fmt.space_after),
        "line_spacing": fmt.line_spacing,
        "line_spacing_rule": None if fmt.line_spacing_rule is None else int(fmt.line_spacing_rule),
        "left_indent_pt": length(fmt.left_indent),
        "right_indent_pt": length(fmt.right_indent),
        "first_line_indent_pt": length(fmt.first_line_indent),
        "keep_with_next": fmt.keep_with_next,
        "keep_together": fmt.keep_together,
        "page_break_before": fmt.page_break_before,
        "runs": runs,
    }


def main():
    source = Path(sys.argv[1])
    out = Path(sys.argv[2])
    doc = Document(source)
    paragraphs = [para_info(i, p) for i, p in enumerate(doc.paragraphs)]
    font_counts = Counter()
    size_counts = Counter()
    for p in paragraphs:
        for run in p["runs"]:
            font_counts[str(run["font"])] += len(run["text"])
            size_counts[str(run["size_pt"])] += len(run["text"])
    sections = []
    for i, sec in enumerate(doc.sections):
        sections.append(
            {
                "index": i,
                "page_width_pt": length(sec.page_width),
                "page_height_pt": length(sec.page_height),
                "margins_pt": {
                    "top": length(sec.top_margin),
                    "bottom": length(sec.bottom_margin),
                    "left": length(sec.left_margin),
                    "right": length(sec.right_margin),
                    "header": length(sec.header_distance),
                    "footer": length(sec.footer_distance),
                },
                "different_first_page": sec.different_first_page_header_footer,
                "header": [p.text for p in sec.header.paragraphs],
                "footer": [p.text for p in sec.footer.paragraphs],
                "first_header": [p.text for p in sec.first_page_header.paragraphs],
                "first_footer": [p.text for p in sec.first_page_footer.paragraphs],
            }
        )
    tables = []
    for i, table in enumerate(doc.tables):
        tables.append(
            {
                "index": i,
                "style": None if table.style is None else table.style.name,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "cell_preview": [[cell.text[:80] for cell in row.cells] for row in table.rows[:3]],
            }
        )
    with zipfile.ZipFile(source) as zf:
        package = []
        for name in sorted(zf.namelist()):
            data = zf.read(name)
            package.append({"name": name, "size": len(data), "sha256": hashlib.sha256(data).hexdigest()})
    data = {
        "source": str(source),
        "sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "font_char_counts": font_counts,
        "size_char_counts": size_counts,
        "sections": sections,
        "tables": tables,
        "inline_shapes": len(doc.inline_shapes),
        "package": package,
    }
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2, default=str), encoding="utf-8")


if __name__ == "__main__":
    main()
