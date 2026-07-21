from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "맑은 고딕"
CODE_FONT = "Consolas"


def set_run_font(run, name=FONT, size=11, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for element in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(element)
    set_run_font(run, size=9)


def keep_lines(paragraph, keep_next=False):
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = keep_next


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(130)
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run(text)
    set_run_font(run, size=30, bold=True)
    return p


def add_chapter(doc, text, new_page=True):
    if new_page and len(doc.paragraphs) > 0:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    keep_lines(p, True)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True)
    return p


def add_section(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    keep_lines(p, True)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.45
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-10)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run("•  " + text)
    set_run_font(run)
    return p


def add_box(doc, title, lines, dark=False):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0)
    cell.width = Inches(6.45)
    set_cell_margins(cell, 130, 160, 130, 160)
    set_cell_shading(cell, "F2F2F2" if not dark else "202124")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=(255, 255, 255) if dark else None)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(line)
        set_run_font(r, name=CODE_FONT if dark else FONT, size=9.5 if dark else 10.5,
                     color=(235, 235, 235) if dark else None)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths:
            cell.width = Inches(widths[i])
        set_cell_shading(cell, "E7E6E6")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if widths:
                cells[i].width = Inches(widths[i])
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            set_run_font(r, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_screenshot_placeholder(doc, number, subject, instruction):
    add_box(
        doc,
        f"[스크린샷 삽입 위치 {number}] {subject}",
        [f"촬영·삽입 안내: {instruction}", "권장: 본문 폭에 맞춰 삽입하고 계정명·토큰·환경변수 등 민감정보는 가린다."],
    )


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Pt(35.4)
    section.footer_distance = Pt(35.4)
    section.different_first_page_header_footer = True
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    for footer in (section.footer,):
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)
        add_page_field(footer.add_paragraph())
    first = section.first_page_footer
    for p in first.paragraphs:
        p.clear()


SECURITY_ITEMS = [
    ("SEC-001", "제재 계정 요청의 미들웨어 순서 오류", "config/settings.py, users/middleware.py", "계정 상태 검사 미들웨어가 메시지 미들웨어보다 먼저 실행되면 제재 안내를 기록하는 시점에 메시지 저장소가 준비되지 않아 서버 오류가 발생할 수 있었다.", "제재 사용자가 인증된 상태로 보호 페이지를 요청한다.", "정상적인 접근 차단 대신 500 오류가 날 수 있고 상세 오류 노출 및 가용성 저하로 이어진다.", "메시지 미들웨어 이후에 상태 검사가 실행되도록 순서를 교정하고 제재 계정의 리다이렉트·메시지 동작을 회귀 테스트했다.", "CWE-20 / OWASP A04:2021"),
    ("SEC-002", "개발 기본 비밀키와 배포 보안 설정 부족", "config/settings.py, .env.example", "초기 개발용 SECRET_KEY가 배포 검사 기준보다 짧고 HTTPS/HSTS 설정의 환경 분리가 충분하지 않았다.", "DEBUG=false 상태에서 Django deploy check를 실행한다.", "예측 가능한 키는 서명 쿠키와 토큰의 신뢰성을 훼손할 수 있고, 안전하지 않은 전송 설정은 세션 보호를 약화한다.", "긴 개발 표식 키를 사용하되 운영에서는 환경변수 난수 키를 필수로 했다. Secure 쿠키·SSL 리다이렉트·HSTS를 환경변수로 분리하고 .env 제외 및 배포 검사를 자동화했다.", "CWE-798 / OWASP A02:2021"),
    ("SEC-003", "로그인 무차별 대입 제한 부재", "security_controls/, users/views.py", "반복 로그인 실패를 제한하는 서버 측 통제가 없었다.", "동일한 사용자명과 IP 조합으로 잘못된 비밀번호를 연속 제출한다.", "공격자가 충분한 횟수의 비밀번호 추측을 수행할 수 있다.", "사용자명과 IP 원문 대신 SECRET_KEY 기반 HMAC 키를 저장하고 실패 창·잠금 만료를 DB에서 관리했다. 성공 시 기록을 삭제하고 잠금·초기화·개인정보 비저장 테스트를 추가했다.", "CWE-307 / OWASP A07:2021"),
    ("SEC-004", "WebSocket 직접 접근과 메시지 도배", "chat/consumers.py, config/asgi.py", "HTTP 화면 권한만으로는 비로그인 소켓 연결, 다른 대화 구독, 변조 Origin 및 고속 메시지 전송을 막을 수 없었다.", "웹 화면을 거치지 않고 WebSocket URL에 연결하거나 다른 사용자의 대화 ID를 지정한다.", "비공개 대화 노출, 저장형 XSS 입력 축적, 저장공간·CPU 고갈 가능성이 있다.", "세션 인증과 AllowedHostsOriginValidator, 연결·수신 시 활성 계정/참여자 재검증, 1000자 제한, 사용자 행 잠금 기반 분당 제한을 적용했다.", "CWE-284·770 / OWASP A01:2021"),
    ("SEC-005", "최초 레코드 생성의 동시성 경쟁", "chat/services.py, security_controls/services.py", "동일한 직접 대화나 로그인 제한 행을 동시에 최초 생성하면 고유 제약 예외가 사용자에게 노출될 수 있었다.", "서로 같은 키로 병렬 요청을 보낸다.", "간헐적 500 오류와 서비스 불안정이 발생한다.", "중첩 savepoint로 충돌을 격리한 뒤 확정된 행을 잠금 재조회하도록 보강하고 경쟁 조건 회귀 테스트를 작성했다.", "CWE-362 / OWASP A04:2021"),
    ("SEC-006", "메시지 관리자 변경의 감사 누락", "chat/admin.py, chat/services.py", "Django Admin에서 메시지 상태를 바꿀 수 있었으나 변경 사유와 전후 값이 감사 로그에 남지 않았다.", "관리자가 메시지를 hidden 또는 deleted로 변경한다.", "관리자 오남용이나 분쟁 발생 시 책임 추적이 어렵다.", "사유를 필수화하고 상태 변경과 불변 AdminAuditLog 생성을 하나의 트랜잭션으로 묶었다. 대화와 멤버는 읽기 전용으로 제한했다.", "CWE-778 / OWASP A09:2021"),
    ("SEC-007", "Django Admin 직접 변경에 의한 감사 서비스 우회", "users/admin.py, products/admin.py, moderation/admin.py", "사용자·상품·신고를 기본 Admin 폼에서 직접 바꾸면 전용 제재 서비스의 권한·사유·감사 로직을 우회할 수 있었다.", "관리자가 모델 변경 URL에 직접 POST한다.", "전후 상태가 없는 제재와 복구가 가능해 관리 통제의 일관성이 깨진다.", "해당 관리 화면을 조회 전용으로 만들고 전용 플랫폼 관리 대시보드에서만 감사되는 서비스 호출을 허용했다. 직접 POST 403과 데이터 불변을 검증했다.", "CWE-284·778 / OWASP A01:2021"),
    ("SEC-008", "불변 기록의 QuerySet 우회 변경", "wallets/models.py, adjustments/models.py, moderation/models.py", "인스턴스 save/delete만 차단하면 QuerySet.update, bulk_update, delete로 원장·제재·감사 기록을 변경할 수 있었다.", "관리 코드 또는 셸에서 불변 모델 QuerySet.update()를 호출한다.", "거래 잔액과 감사 증거의 무결성이 훼손될 수 있다.", "불변 전용 QuerySet/Manager를 적용해 일괄 수정·삭제도 모델 계층에서 거절하고 네 종류 기록에 회귀 테스트를 추가했다.", "CWE-664 / OWASP A08:2021"),
    ("SEC-009", "서로 다른 대상을 이용한 신고 도배", "moderation/services.py", "동일 대상 중복만 차단하면 공격자가 여러 상품·사용자를 번갈아 신고할 수 있었다.", "한 계정으로 제한 시간 안에 서로 다른 대상을 연속 신고한다.", "관리자 업무 방해와 자동 제재 임계치 악용이 가능하다.", "신고자 행을 잠그고 활성 상태를 재검증한 뒤 시간당 사용자별 상한을 적용했다. 한도 초과와 비활성 신고자의 차단을 테스트했다.", "CWE-770 / OWASP A04:2021"),
    ("SEC-010", "CSP·Permissions-Policy 부재와 인라인 스크립트", "config/middleware.py, templates/, static/", "기본 보안 헤더에 CSP와 Permissions-Policy가 없고 채팅 코드가 템플릿 인라인 JavaScript로 존재했다.", "응답 헤더와 템플릿의 script/style 속성을 검사한다.", "XSS가 발생했을 때 실행 범위가 넓고 카메라·마이크 등 불필요한 브라우저 기능이 열릴 수 있다.", "JS/CSS를 정적 파일로 분리하고 script-src/style-src unsafe-inline을 제거했다. object-src none, frame-ancestors none, Permissions-Policy와 Bootstrap SRI를 적용했다.", "CWE-693 / OWASP A05:2021"),
    ("SEC-011", "가격·송금·조정 금액의 업무 상한 부재", "products/models.py, wallets/services.py, adjustments/services.py", "DB 정수 범위만 의존해 비정상적으로 큰 가격과 포인트 이동이 가능했다.", "정수형 한계에 가까운 가격·송금·조정 요청을 보낸다.", "잔액 모델의 의미가 깨지고 서비스 거부 또는 운영 실수가 확대될 수 있다.", "상품 가격 DB 제약과 폼 검증, 1회 송금·조정 상한, 지갑 최대 잔액을 서비스 계층에 추가하고 경계값·초과 요청을 테스트했다.", "CWE-20 / OWASP A04:2021"),
    ("SEC-012", "DEBUG=false에서 채팅 정적 파일 누락", "Dockerfile, docker-compose.yml, scripts/run_e2e.sh", "CSP 대응으로 채팅 JS를 정적 파일화했지만 운영형 시작 절차가 collectstatic을 실행하지 않아 채팅 버튼이 동작하지 않았다.", "DEBUG=false Docker/E2E 환경에서 1대1 채팅 전송을 실행한다.", "보안 강화가 핵심 기능 가용성을 깨뜨리는 회귀가 발생했다.", "이미지 빌드·Compose·E2E 시작에 collectstatic을 추가하고 finders 검사, 이미지 내부 파일 존재, Playwright 전체 흐름으로 검증했다.", "CWE-16 / OWASP A05:2021"),
]


def build(template, output):
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)
    doc = Document(output)
    clear_body(doc)
    configure(doc)
    props = doc.core_properties
    props.title = "Tiny Second-hand Shopping Platform 개발 및 보안 검증 보고서"
    props.subject = "WhiteHat School 시큐어 코딩 과제"
    props.author = "[이름]"
    props.comments = "스크린샷은 사용자 삽입용 위치 안내만 포함"

    add_title(doc, "Tiny Second-hand Shopping Platform\n개발 및 보안 검증 보고서")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("[분반 / 이름 / 전화번호 뒤 4자리]"), size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("작성일 ")
    set_run_font(r, bold=True)
    set_run_font(p.add_run("2026년 7월 20일"))

    add_chapter(doc, "1. 프로젝트 개요")
    add_section(doc, "1.1 개발 목적과 범위")
    add_body(doc, "본 프로젝트는 WhiteHat School 시큐어 코딩 과제의 최소 요구사항을 실제 데이터베이스, 서버 측 인증·권한 검증, 입력 검증, 실패 처리와 자동화 테스트까지 연결해 완성한 교육용 중고거래 플랫폼이다. 단순 화면 시연이 아니라 회원·상품·채팅·신고·제재·포인트 송금·관리자 감사 기능이 서로 일관된 상태를 유지하도록 설계했다.")
    add_body(doc, "실물 화폐나 외부 결제는 다루지 않고 플랫폼 내부 테스트 포인트만 사용한다. 실행 범위는 로컬 Docker Compose와 GitHub Actions까지이며, Django 서버 렌더링 모놀리스로 프런트엔드 복잡도를 줄이고 보안 통제를 서버에 집중시켰다.")
    add_box(doc, "프로젝트 식별 정보", ["저장소: https://github.com/choi95411/secure-coding", "실행 형태: Django + PostgreSQL + Redis + Docker Compose", "검증 기준: 과제 PDF 21~36쪽 및 35쪽 최소 기능", "최종 로컬 회귀: 68 passed, 2 skipped", "최종 CI: GitHub Actions run 29724174781 — 전체 통과"])
    add_section(doc, "1.2 개발 원칙")
    for item in ["모든 접근 제어는 화면 표시 여부가 아니라 서버의 인증·소유권·참여자·관리자 검사로 결정한다.", "사용자 입력은 Django Form/Model/Service 경계에서 검증하고 SQL은 Django ORM의 파라미터 바인딩을 사용한다.", "송금과 관리자 조정은 트랜잭션·행 잠금·불변 원장·멱등 키로 무결성을 보장한다.", "기능마다 정상·실패·비로그인·권한 위반 테스트를 작성하고, 마일스톤마다 회귀·문서·Git 기록을 갱신한다.", "운영 비밀은 환경변수로만 주입하며 .env, 토큰, 개인키와 개인정보를 Git에 커밋하지 않는다."]:
        add_bullet(doc, item)

    add_chapter(doc, "2. 요구사항 분석")
    add_section(doc, "2.1 필수 기능 해석")
    add_body(doc, "과제의 열 가지 최소 기능을 화면 단위가 아니라 업무 규칙 단위로 분해했다. 각 요구사항에는 사용자 역할, 선행 조건, 정상·실패 흐름, 인증·권한 조건, 보안 조건, 인수 조건, 구현 파일과 테스트를 연결했다. 상세 추적 정보는 docs/requirements.md에 유지한다.")
    req_rows = [
        ("R1", "회원가입·로그인·로그아웃", "세션 인증, 비밀번호 해시, 무차별 대입 제한", "완료"),
        ("R2", "사용자 조회·프로필·마이페이지", "활성 상태 필터, 본인 변경", "완료"),
        ("R3", "상품 CRUD·이미지", "소유자 검사, 업로드 검증", "완료"),
        ("R4", "전체·1대1 채팅", "참여자·Origin·Rate Limit", "완료"),
        ("R5", "사용자·상품 신고", "중복·시간당 상한", "완료"),
        ("R6", "악성 상품 차단", "임계치 자동 차단·관리자 복구", "완료"),
        ("R7", "악성 사용자 휴면·차단", "상태 미들웨어·감사", "완료"),
        ("R8", "사용자 간 송금", "트랜잭션·잠금·멱등·원장", "완료"),
        ("R9", "상품 검색", "공개 범위·상태·정렬·페이지", "완료"),
        ("R10", "관리자 통합 관리", "서버 권한·사유·전후 값·감사", "완료"),
    ]
    add_table(doc, ["ID", "기능", "핵심 인수 조건", "상태"], req_rows, [0.5, 1.7, 3.6, 0.65])
    add_section(doc, "2.2 정상 흐름과 실패 흐름")
    add_body(doc, "정상 흐름은 가입 후 로그인, 상품 등록, 다른 사용자의 검색·상세 조회, 1대1 대화, 포인트 송금, 신고와 관리자 제재로 이어진다. 실패 흐름은 빈 값·형식 오류·중복 요청·없는 대상·비로그인·타인 데이터·정지 계정·동시 요청으로 구분했다. 특히 송금·신고·관리자 기능은 성공 화면보다 실패 시 데이터가 변하지 않는지를 인수 조건으로 삼았다.")
    add_section(doc, "2.3 비기능 요구사항")
    for item in ["보안성: SQL Injection, XSS, CSRF, IDOR, 권한 상승, 파일 업로드, WebSocket, 송금 경쟁 조건을 회귀 테스트한다.", "실행 가능성: README 절차로 Docker/비 Docker 실행, 마이그레이션, 관리자·테스트 데이터 생성을 재현한다.", "감사 가능성: 원장·제재·관리자 작업은 삭제·수정이 불가능하고 사유와 전후 상태를 남긴다.", "유지보수성: 앱별 책임 분리, docs/progress.md 영속 상태, 작은 Git 커밋과 CI를 사용한다."]:
        add_bullet(doc, item)

    add_chapter(doc, "3. 시스템 설계")
    add_section(doc, "3.1 전체 구조")
    add_body(doc, "시스템은 Django 서버 렌더링 모놀리스다. 브라우저의 HTTP 요청은 URL·View·Form을 거쳐 서비스와 ORM으로 전달되고, 템플릿은 autoescape를 유지한다. WebSocket 요청은 ASGI의 Origin 검증과 세션 인증을 통과한 뒤 소비자에서 계정 상태와 대화 참여자를 다시 확인한다. PostgreSQL이 모든 영구 업무 데이터의 원본이며 Redis는 Channels의 프로세스 간 전달에만 사용한다.")
    add_box(doc, "요청 처리 경계", ["HTTP: Browser → Django URL/View/Form → Service → ORM → PostgreSQL", "WebSocket: Origin 검사 → 세션 인증 → 참여자/상태 검사 → Message 저장 → Redis broadcast", "관리: 관리자 권한 → 사유 검증 → 전용 서비스 → 대상 변경 + AdminAuditLog", "송금: 입력 검증 → atomic → Wallet 행 잠금 → Transfer + LedgerEntry"])
    add_section(doc, "3.2 애플리케이션 책임")
    app_rows = [("users", "인증, 프로필, 계정 상태"), ("products", "상품 CRUD, 이미지, 공개 검색"), ("wallets", "지갑, 송금, 불변 원장"), ("adjustments", "관리자 포인트 조정"), ("moderation", "신고, 제재, 감사 로그"), ("chat", "전체·1대1 대화, 메시지, WebSocket"), ("security_controls", "개인정보 비저장 로그인 제한"), ("config", "설정, 보안 헤더, ASGI/WSGI")]
    add_table(doc, ["모듈", "책임"], app_rows, [1.55, 4.9])
    add_section(doc, "3.3 권한 모델")
    add_table(doc, ["대상", "비로그인", "일반 사용자", "소유자/참여자", "관리자"], [
        ("공개 상품", "조회", "조회·검색", "수정·삭제", "전체 조회·제재"),
        ("프로필", "활성 공개 정보", "활성 공개 정보", "본인 변경", "조회·상태 관리"),
        ("1대1 채팅", "차단", "차단", "열람·전송", "정책 범위 관리"),
        ("송금/원장", "차단", "본인 송금·내역", "본인 내역", "조회·조정 거래"),
        ("관리 화면", "차단", "403", "403", "사유 기반 작업"),
    ], [1.25, 1.05, 1.6, 1.55, 1.5])

    add_chapter(doc, "4. 데이터베이스 및 핵심 무결성 설계")
    add_section(doc, "4.1 주요 모델과 관계")
    add_body(doc, "User와 UserProfile, Wallet은 1:1이다. Product는 판매자를 PROTECT로 참조하고 ProductImage를 종속시킨다. Conversation과 User는 ConversationMember로 연결되며 Message는 대화와 발신자를 PROTECT로 보존한다. Report·ModerationAction·AdminAuditLog는 신고와 관리자 조치의 증거를 유지한다. Transfer는 송금 업무 상태를, LedgerEntry는 차변·대변 원장을 나타낸다.")
    model_rows = [
        ("User/UserProfile", "고유 username, active/dormant/blocked", "삭제 대신 상태 변경"),
        ("Product/Image", "공개·판매·차단·삭제 상태, 가격 제약", "UUID 경로·검증 이미지"),
        ("Conversation/Member", "global 조건부 단일, direct_key 고유", "참여자별 접근"),
        ("Message", "1000자, visible/hidden/deleted", "발신자·대화 PROTECT"),
        ("Wallet/Transfer/Ledger", "멱등 키, 행 잠금, debit/credit", "원장 수정·삭제 금지"),
        ("Report/Moderation/Audit", "중복 신고, 사유, 전후 JSON", "감사 기록 불변"),
    ]
    add_table(doc, ["모델군", "제약·상태", "보존 정책"], model_rows, [1.75, 2.9, 1.8])
    add_section(doc, "4.2 포인트 송금의 원자성")
    add_body(doc, "송금은 양의 정수, 자기 자신 금지, 활성 수신자, 1회 상한, 잔액 상한을 먼저 검증한다. 하나의 DB 트랜잭션에서 두 지갑을 결정적 순서로 select_for_update 잠금하고 잔액을 확인한 다음 송금·양쪽 원장을 기록한다. 같은 멱등 키는 동일 결과를 반환하며 실패 시 어떤 잔액과 원장도 남지 않는다.")
    add_box(doc, "송금 불변식", ["sender.balance ≥ 0", "송금 전 총잔액 = 송금 후 총잔액", "성공 Transfer 1건 ↔ debit/credit LedgerEntry 각 1건", "동일 idempotency key는 효과가 한 번만 발생", "관리자는 잔액 직접 수정 대신 조정 거래 생성"])
    add_section(doc, "4.3 검색과 공개 범위")
    add_body(doc, "상품명과 설명을 ORM icontains 조건으로 검색하고 공개·판매 상태만 노출한다. 빈 검색어는 공개 전체 목록으로 처리하며 정렬·페이지네이션을 서버에서 적용한다. 비공개·삭제·차단 상품은 일반 사용자 쿼리셋 자체에서 제외하므로 URL 직접 접근에도 노출되지 않는다. 관리 목적 조회는 관리자 전용 엔드포인트에서만 허용한다.")

    add_chapter(doc, "5. 구현 과정")
    add_section(doc, "5.1 기반·인증·상품·검색")
    add_body(doc, "초기 환경과 과제 PDF를 조사한 뒤 Django 프로젝트, 환경변수, PostgreSQL 마이그레이션과 Bootstrap 기반 공통 화면을 구성했다. Django 내장 인증·세션·비밀번호 해시·CSRF를 재사용하고 가입 폼에 강도와 중복 검증을 적용했다. 프로필과 마이페이지는 본인 변경만 허용하며, 계정 상태 미들웨어가 휴면·차단 사용자의 보호 기능 접근을 막는다.")
    add_body(doc, "상품 생성·수정·삭제는 로그인과 소유자 검사를 모두 수행한다. 이미지 업로드는 확장자, MIME, 크기와 실제 이미지 디코딩을 검증하고 UUID 경로를 사용한다. 검색은 raw SQL 없이 ORM으로 구현하고 상태 필터와 페이지네이션을 테스트했다.")
    add_section(doc, "5.2 지갑·송금·원장")
    add_body(doc, "신규 사용자는 정책에 따른 초기 테스트 포인트를 지갑으로 지급받는다. 송금 서비스는 두 지갑 잠금, 멱등 키, 실패 롤백과 원장 생성까지 한 트랜잭션으로 처리한다. PostgreSQL 외부 통합 테스트에서는 동시에 잔액을 초과하는 송금을 시도해 하나만 성공하고 음수 잔액이 발생하지 않음을 확인했다. 관리자 조정도 별도 Adjustment와 LedgerEntry, 감사 로그를 함께 생성한다.")
    add_section(doc, "5.3 신고·제재·관리자")
    add_body(doc, "사용자와 상품 신고는 동일 대상 중복뿐 아니라 사용자별 시간당 상한을 적용한다. 신고 누적 임계치에 도달하면 상품은 차단되고 사용자는 휴면 또는 차단 상태로 전환된다. 관리자는 전용 대시보드에서 전체 사용자·상품·신고·거래를 조회하고 사유를 입력해 제재·복구·메시지 상태 변경·포인트 조정을 수행한다. 모든 변경은 관리자, 대상, 작업, 사유, 전후 값과 시간을 불변 감사 로그에 기록한다.")
    add_section(doc, "5.4 전체·1대1 채팅")
    add_body(doc, "전체 대화는 조건부 단일 제약으로 하나만 유지하고, 1대1 대화는 정렬한 두 사용자 ID로 direct_key를 생성해 중복을 막는다. 메시지는 PostgreSQL에 먼저 저장한 뒤 Redis channel layer로 전달한다. WebSocket 연결과 메시지 수신 때마다 세션, Origin, 활성 계정과 참여자 조건을 확인하며 길이와 분당 전송량을 제한한다. 화면 출력은 템플릿 autoescape와 DOM textContent를 사용한다.")
    add_section(doc, "5.5 관리자와 사용자 화면")
    add_body(doc, "화면은 서버 렌더링과 Bootstrap을 사용해 가입·로그인·상품·채팅·지갑·신고·관리 흐름을 일관되게 구성했다. 보안 강화 과정에서 인라인 JavaScript와 CSS를 정적 파일로 이동했고, 403·404·500 페이지는 내부 예외 정보를 노출하지 않으면서 사용자가 다음 행동을 이해하도록 작성했다.")

    add_chapter(doc, "6. 보안 설계와 위협 대응")
    add_section(doc, "6.1 인증·세션·접근 제어")
    add_body(doc, "Django의 PBKDF2 계열 비밀번호 해시, 세션 순환, CSRF 미들웨어와 HttpOnly·SameSite 쿠키를 기본으로 한다. 로그인 실패는 HMAC 처리된 사용자명/IP 결합 키로 제한해 원문 개인정보를 저장하지 않는다. 객체 접근은 소유자·참여자·관리자 조건을 View와 WebSocket 소비자에서 검사하고 권한 위반은 403 또는 4403으로 종료한다.")
    add_section(doc, "6.2 입력·출력·파일 보안")
    add_body(doc, "모든 DB 검색과 변경은 ORM을 사용해 SQL 파라미터 바인딩을 유지한다. 템플릿 autoescape와 textContent로 저장형·반사형·DOM XSS를 막으며 CSP를 추가 방어선으로 둔다. 상태 변경은 POST와 CSRF 토큰을 요구한다. 이미지 파일은 확장자·MIME·크기·내용을 함께 검증하고 사용자가 제공한 파일명을 저장 경로에 사용하지 않는다.")
    add_section(doc, "6.3 관리자·감사·불변성")
    add_body(doc, "관리자 여부는 서버의 is_staff 조건으로 검증하며, 버튼 숨김은 편의 기능일 뿐 권한 통제가 아니다. 사용자·상품·신고·원장·조정·감사 모델의 Django Admin 직접 변경을 막아 전용 감사 서비스를 강제한다. 불변 QuerySet은 update, bulk_update, delete를 거절해 셸이나 관리 코드도 원장과 감사 증거를 변경할 수 없게 한다.")
    add_section(doc, "6.4 보안 헤더와 배포")
    add_body(doc, "X-Frame-Options, nosniff, CSP, Permissions-Policy를 적용하고 Bootstrap CDN 자원에는 SRI를 사용한다. 운영 설정은 DEBUG=false, 난수 SECRET_KEY, 허용 호스트, Secure 쿠키와 HSTS를 환경변수로 제어한다. 컨테이너는 appuser 비루트 계정으로 실행하고 정적 파일은 이미지 빌드와 시작 절차에서 수집한다.")

    add_chapter(doc, "7. 체크리스트 및 테스트")
    add_section(doc, "7.1 테스트 계층")
    test_rows = [
        ("모델/서비스", "제약, 상태 전이, 불변 QuerySet, 원장 정합성"),
        ("인증/권한", "가입·로그인·잠금, 비로그인, 소유권, 관리자 403"),
        ("통합", "PostgreSQL 송금 동시성, Redis channel layer"),
        ("보안 회귀", "SQLi 문자열, XSS, CSRF, IDOR, 업로드, 헤더"),
        ("WebSocket", "미인증, Origin, 비참여자, 길이·Rate Limit"),
        ("Playwright E2E", "A/B 가입→상품→검색→채팅→송금→신고→제재"),
    ]
    add_table(doc, ["계층", "검증 범위"], test_rows, [1.7, 4.75])
    add_section(doc, "7.2 핵심 시나리오")
    for item in ["사용자 A와 B 가입, A 상품 등록, B 검색·상세 조회", "A와 B의 1대1 채팅 및 WebSocket 메시지 영속화", "B의 테스트 포인트 송금과 양쪽 잔액·원장 확인", "상품 또는 사용자 신고 후 관리자의 사유 기반 제재", "일반 사용자의 관리자 URL/API 직접 호출 403", "동일 멱등 키 중복 송금 1회 처리 및 동시 잔액 초과 방지"]:
        add_bullet(doc, item)
    add_section(doc, "7.3 최종 실행 결과")
    add_box(doc, "최종 검증 결과", ["pytest: 68 passed, 2 skipped (로컬 SQLite에서 PostgreSQL/Redis 전용 테스트만 조건부 skip)", "GitHub Actions: PostgreSQL 17·Redis 7.4 전체 테스트 통과", "Ruff lint·format check·migration drift·Django check 통과", "Bandit·pip check·pip-audit·deploy check 통과", "Git 전체 이력 .env·토큰·개인키 검사 통과", "Docker 29.2.1 빌드, appuser 비루트, 컨테이너 Django check 통과", "Playwright 핵심 E2E 및 증빙 업로드 통과", "최종 CI: https://github.com/choi95411/secure-coding/actions/runs/29724174781"])
    add_section(doc, "7.4 보안 체크리스트 판정")
    checklist = ["SQL Injection: ORM/파라미터 바인딩 및 공격 문자열 회귀 통과", "Stored/Reflected/DOM XSS: autoescape·textContent·CSP 회귀 통과", "CSRF·세션·쿠키: 미들웨어와 운영 설정 검사 통과", "IDOR·Broken Access Control: 상품 소유자·채팅 참여자·관리자 403 통과", "업로드: 확장자·MIME·크기·실제 이미지·UUID 경로 통과", "송금: 양수·자기송금·수신자·잔액·멱등·롤백·동시성 통과", "WebSocket: 인증·Origin·참여자·길이·Rate Limit 통과", "비밀정보·의존성·Debug: 이력 검사·pip-audit·deploy check 통과"]
    for item in checklist:
        add_bullet(doc, item)

    add_chapter(doc, "8. 개발 과정에서 발견한 보안 약점과 수정")
    add_body(doc, "다음 항목은 설계·구현·회귀·CI 과정에서 실제로 확인한 약점이다. 각 수정은 기능을 숨기는 수준이 아니라 서버·모델·배포 경계의 통제로 반영했고 재현 조건에 대응하는 자동화 테스트를 남겼다.")
    for idx, item in enumerate(SECURITY_ITEMS, start=1):
        sec_id, title, location, description, reproduce, impact, fix, mapping = item
        add_section(doc, f"8.{idx} {sec_id} — {title}")
        add_body(doc, f"발견 위치: {location}", bold_lead="발견 위치:")
        add_body(doc, f"약점 설명: {description}", bold_lead="약점 설명:")
        add_body(doc, f"재현 및 수정 전 동작: {reproduce} 수정 전에는 {impact}", bold_lead="재현 및 수정 전 동작:")
        add_body(doc, f"수정 내용과 검증: {fix}", bold_lead="수정 내용과 검증:")
        add_body(doc, f"분류: {mapping}", bold_lead="분류:")

    add_chapter(doc, "9. 유지보수·DevOps·GitHub 관리")
    add_section(doc, "9.1 반복 개발과 회귀 수정")
    add_body(doc, "각 기능은 요구사항 확인, 위협 검토, 구현, 관련 테스트, 코드 검토, 문서 갱신, Git 커밋 순서로 진행했다. PostgreSQL 동시성 테스트에서 스레드 연결이 닫히지 않아 테스트 DB 삭제가 실패한 문제는 스레드별 연결 종료로 수정했다. 관리자 읽기 전용 공통화가 감사되는 메시지 변경까지 막은 회귀는 예외 흐름을 복원하고 관련 테스트로 고정했다.")
    add_body(doc, "첫 보안 강화 CI에서는 DEBUG=false에서 정적 채팅 파일이 수집되지 않아 Playwright 전송 버튼 단언이 실패했다. collectstatic을 Dockerfile·Compose·E2E 시작 과정에 추가하고 이미지 내부 파일, Django finders, 전체 브라우저 흐름을 확인했다. 이 사례는 보안 변경 뒤 기능 가용성까지 통합 검증해야 함을 보여준다.")
    add_section(doc, "9.2 CI와 공급망 검사")
    add_body(doc, "GitHub Actions는 PostgreSQL·Redis 서비스, 전체 pytest, Ruff, Bandit, pip-audit, migration drift, deploy check, 비밀 이력 검사, Docker 빌드와 Playwright E2E를 실행한다. Dependabot은 pip와 Actions 의존성을 주간 점검하지만 자동 병합하지 않고 검토 후 반영한다. checkout, setup-python, upload-artifact는 당시 공식 최신 주요 버전인 v7로 갱신했다.")
    add_section(doc, "9.3 문서와 재현성")
    add_body(doc, "README에는 Docker/비 Docker 설치, 환경변수, 마이그레이션, 관리자와 테스트 데이터, 실행·테스트 명령, 보안 설계와 교육용 포인트 제한을 기록했다. docs/requirements.md는 요구사항-구현-테스트 추적표, docs/security-checklist.md는 약점 수정 기록, docs/progress.md는 재개 가능한 최신 상태를 제공한다.")

    add_chapter(doc, "10. 완성도 평가와 결론")
    add_section(doc, "10.1 최소 요구사항 충족")
    add_body(doc, "회원가입·로그인·로그아웃, 사용자 조회·프로필, 상품 CRUD·검색, 전체·1대1 채팅, 사용자·상품 신고, 악성 상품 차단, 악성 사용자 휴면·차단, 포인트 송금, 관리자 통합 관리가 모두 DB와 연결되어 동작한다. 각 기능은 정상·실패·비로그인·권한 위반 테스트에 연결되며 최종 CI에서 실제 PostgreSQL·Redis·브라우저 흐름까지 통과했다.")
    add_section(doc, "10.2 알려진 제한사항")
    for item in ["교육용 내부 포인트만 사용하며 은행·카드·실물 화폐 결제와 연결하지 않는다.", "배포 대상은 로컬 Docker와 GitHub Actions 검증까지이며 공개 운영 인프라, 도메인, TLS 종료는 범위 밖이다.", "로컬 SQLite 실행에서는 PostgreSQL 행 잠금 및 Redis 외부 통합 테스트 2건이 skip되지만 CI에서는 실제 서비스로 실행된다.", "관리자의 채팅 관리 범위는 메시지 상태 관리와 감사로 제한하며 비공개 대화의 무제한 열람을 제공하지 않는다.", "스크린샷은 개인정보와 실행 환경 차이를 고려해 본 문서에는 넣지 않고 사용자가 최종 제출본에 직접 삽입한다."]:
        add_bullet(doc, item)
    add_section(doc, "10.3 결론")
    add_body(doc, "이 프로젝트는 필수 기능의 존재뿐 아니라 데이터 무결성, 서버 측 권한, 공격 입력, 동시 요청, 관리자 책임 추적과 새 환경 재현성을 완료 조건으로 삼았다. 개발 중 발견된 열두 가지 약점을 즉시 문서화하고 회귀 테스트로 고정했으며, 최종적으로 코드·보안·Docker·PostgreSQL·Redis·Playwright 검증이 하나의 CI에서 통과하는 상태를 만들었다.")

    add_chapter(doc, "부록 A. 스크린샷 삽입 위치 안내")
    add_body(doc, "아래에는 실제 이미지를 넣지 않고 최종 제출본에서 사용자가 직접 촬영·삽입할 위치와 내용을 표시한다. 계정명, 세션, 토큰, 환경변수와 로컬 경로 등 민감정보가 보이지 않도록 확인한다.")
    add_screenshot_placeholder(doc, 1, "메인 상품 목록과 검색", "로그인 후 공개 상품 목록에서 검색어·정렬·페이지네이션과 상품 카드가 함께 보이도록 촬영한다.")
    add_screenshot_placeholder(doc, 2, "상품 등록 및 상세", "사용자 A가 상품을 등록한 결과와 사용자 B가 상세를 조회하는 화면을 각각 촬영한다.")
    add_screenshot_placeholder(doc, 3, "전체 채팅과 1대1 채팅", "서로 다른 브라우저 세션에서 A와 B의 메시지가 실시간 반영되는 화면을 촬영한다.")
    add_screenshot_placeholder(doc, 4, "포인트 송금과 원장", "송금 전후 양쪽 잔액 및 거래 내역이 일치하는 화면을 촬영한다.")
    add_screenshot_placeholder(doc, 5, "신고와 관리자 제재", "신고 상세, 사유 입력, 제재 후 상품 또는 사용자 상태와 감사 로그를 촬영한다.")
    add_screenshot_placeholder(doc, 6, "일반 사용자 관리자 접근 차단", "일반 계정으로 관리자 URL 직접 접근 시 403 화면을 촬영한다.")
    add_screenshot_placeholder(doc, 7, "GitHub Actions 성공", "run 29724174781의 전체 작업 성공 상태와 E2E 증빙 아티팩트를 촬영한다.")

    add_chapter(doc, "부록 B. 실행 및 검증 명령")
    add_box(doc, "Docker 실행", ["cp .env.example .env", "docker compose up --build", "docker compose exec web python manage.py migrate"], dark=True)
    add_box(doc, "로컬 테스트·보안 검사", [".venv/bin/pytest", ".venv/bin/ruff check .", ".venv/bin/ruff format --check .", ".venv/bin/bandit -c pyproject.toml -r config users products wallets moderation adjustments chat security_controls", ".venv/bin/pip-audit -r requirements.txt", "bash scripts/scan_secrets.sh"], dark=True)

    doc.save(output)


if __name__ == "__main__":
    build(Path(sys.argv[1]), Path(sys.argv[2]))
